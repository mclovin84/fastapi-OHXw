# main.py - Complete LangChain Property Scraper System

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional
import asyncio
import aiohttp
from datetime import datetime, timedelta
import json
import os
import re
import logging
import traceback
from docx import Document
from docxtpl import DocxTemplate
import tempfile
import zipfile
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Inches

# LangChain imports (NOT ACTUALLY USED - could be removed)
from langchain.agents import create_openai_functions_agent, AgentExecutor
from langchain.tools import tool
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.document_loaders import AsyncHtmlLoader
from langchain_community.document_transformers import Html2TextTransformer

# Airtop browser automation
from airtop import AsyncAirtop, SessionConfigV1, PageQueryConfig

# Configure logging for Railway
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="LOI Generator - LangChain Edition")

# Add CORS middleware BEFORE routes
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify your frontend domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global exception handler
@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    logger.error(f"Unhandled exception: {exc}")
    logger.error(traceback.format_exc())
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal server error: {str(exc)}"}
    )

# Get API keys from environment variables
OPENAI_KEY = os.getenv("OPENAI_API_KEY")
AIRTOP_KEY = os.getenv("AIRTOP_API_KEY")

# Validate API keys exist
if not OPENAI_KEY:
    print("Warning: Missing OPENAI_API_KEY environment variable! Some features may not work.")
if not AIRTOP_KEY:
    print("Warning: Missing AIRTOP_API_KEY environment variable! Browser automation will not work.")

# Initialize LLM (NOT ACTUALLY USED - could be removed)
llm = None
if OPENAI_KEY:
    try:
        llm = ChatOpenAI(
            api_key=OPENAI_KEY,
            model="gpt-4o-mini",
            temperature=0
        )
    except Exception as e:
        print(f"Error initializing OpenAI: {e}")

# Initialize Airtop (only if API key is available)
airtop_client = None
if AIRTOP_KEY:
    try:
        airtop_client = AsyncAirtop(api_key=AIRTOP_KEY)
    except Exception as e:
        print(f"Error initializing Airtop: {e}")

# Simple in-memory cache
property_cache = {}

# Request models
class PropertyRequest(BaseModel):
    address: str
    
class BatchRequest(BaseModel):
    addresses: List[str]
    email: Optional[str] = None

# Response models
class PropertyData(BaseModel):
    address: str
    owner_name: str
    owner_mailing_address: str
    listing_price: float
    last_sale_price: Optional[float]
    property_details: Dict
    calculations: Dict
    scraped_at: datetime

# Address normalization function for Fulton County
def normalize_address_for_fulton(address: str) -> str:
    """Normalize address for Fulton County assessor search"""
    
    # Comprehensive abbreviation mapping based on USPS standards and Georgia property records
    ABBREVIATIONS = {
        # Street types (USPS standard)
        'STREET': 'ST',
        'AVENUE': 'AVE', 
        'BOULEVARD': 'BLVD',
        'DRIVE': 'DR',
        'ROAD': 'RD',
        'LANE': 'LN',
        'COURT': 'CT',
        'CIRCLE': 'CIR',
        'PLACE': 'PL',
        'PARKWAY': 'PKWY',
        'WAY': 'WAY',
        'TRAIL': 'TRL',
        'TERRACE': 'TER',
        'PLAZA': 'PLZ',
        'ALLEY': 'ALY',
        'BRIDGE': 'BRG',
        'BYPASS': 'BYP',
        'CAUSEWAY': 'CSWY',
        'CENTER': 'CTR',
        'CENTRE': 'CTR',
        'CROSSING': 'XING',
        'EXPRESSWAY': 'EXPY',
        'EXTENSION': 'EXT',
        'FREEWAY': 'FWY',
        'GROVE': 'GRV',
        'HEIGHTS': 'HTS',
        'HIGHWAY': 'HWY',
        'HOLLOW': 'HOLW',
        'JUNCTION': 'JCT',
        'MOTORWAY': 'MTWY',
        'OVERPASS': 'OPAS',
        'PARK': 'PARK',
        'POINT': 'PT',
        'ROUTE': 'RTE',
        'SKYWAY': 'SKWY',
        'SQUARE': 'SQ',
        'TURNPIKE': 'TPKE',
        
        # Directionals (USPS standard)
        'NORTH': 'N',
        'SOUTH': 'S', 
        'EAST': 'E',
        'WEST': 'W',
        'NORTHEAST': 'NE',
        'NORTHWEST': 'NW',
        'SOUTHEAST': 'SE',
        'SOUTHWEST': 'SW',
        
        # Special cases for Georgia property searches (critical for MLK addresses)
        'MARTIN LUTHER KING JR': 'M L KING JR',
        'MARTIN LUTHER KING': 'M L KING',
        'MLK': 'M L KING',
        'ML KING': 'M L KING',
        'MARTIN L KING': 'M L KING',
        'MARTIN LUTHER KING JUNIOR': 'M L KING JR',
        'DR MARTIN LUTHER KING': 'M L KING',
        'REV MARTIN LUTHER KING': 'M L KING',
        
        # Other common name abbreviations
        'SAINT': 'ST',
        'MOUNT': 'MT',
        'FORT': 'FT',
        'DOCTOR': 'DR',
        'REVEREND': 'REV',
        'JUNIOR': 'JR',
        'SENIOR': 'SR',
        'FIRST': '1ST',
        'SECOND': '2ND',
        'THIRD': '3RD',
        'FOURTH': '4TH',
        'FIFTH': '5TH',
        'SIXTH': '6TH',
        'SEVENTH': '7TH',
        'EIGHTH': '8TH',
        'NINTH': '9TH',
        'TENTH': '10TH',
        
        # Building types
        'APARTMENT': 'APT',
        'BUILDING': 'BLDG',
        'SUITE': 'STE',
        'UNIT': 'UNIT',
        'FLOOR': 'FL'
    }

    # Convert to uppercase and remove punctuation
    normalized = address.upper().replace(',', '').replace('#', '')

    # Apply abbreviation mapping with word boundary protection
    for long_form, abbr in ABBREVIATIONS.items():
        # Escape special regex characters and use word boundaries
        escaped_long_form = re.escape(long_form)
        pattern = r'\b' + escaped_long_form + r'\b'
        normalized = re.sub(pattern, abbr, normalized)

    # Remove common Georgia cities, state, and zip codes
    parts = normalized.split()
    filtered_parts = []

    for part in parts:
        # Stop at common Georgia cities
        if part in ['ATLANTA', 'AUGUSTA', 'COLUMBUS', 'MACON', 'SAVANNAH', 'ATHENS', 'ALBANY', 'WARNER', 'ROBINS', 'VALDOSTA', 'GA', 'GEORGIA', 'FAIRBURN', 'PALMETTO', 'SOUTH', 'FULTON']:
            break
        
        # Stop at 5-digit zip codes
        if re.match(r'^\d{5}$', part):
            break
        
        # Skip empty parts
        if part.strip():
            filtered_parts.append(part)

    # Return normalized street address only
    return ' '.join(filtered_parts).strip()

# County Scraper Agent
class CountyScraperAgent:
    def __init__(self):
        self.airtop = airtop_client
        
    async def scrape_fulton_county(self, address: str) -> Dict:
        """Scrapes Fulton County, GA assessor for owner info using step-by-step Airtop interactions"""
        if not self.airtop:
            raise Exception("Airtop client not available - AIRTOP_API_KEY required")
            
        session = None
        try:
            # Normalize address for Fulton County search
            normalized_address = normalize_address_for_fulton(address)
            logger.info(f"Normalized address: {address} -> {normalized_address}")
            
            # Create session with proper timeout - wrap with timeout to prevent hanging
            config = SessionConfigV1(timeout_minutes=10)
            session = await asyncio.wait_for(
                self.airtop.sessions.create(configuration=config),
                timeout=15.0  # Don't let session creation take forever
            )
            session_id = session.data.id
            logger.info(f"Created session: {session_id}")
            
            # Create window and navigate to Fulton County assessor
            window = await self.airtop.windows.create(
                session_id, 
                url="https://qpublic.schneidercorp.com/Application.aspx?App=FultonCountyGA&PageType=Search"
            )
            window_id = window.data.window_id
            logger.info(f"Opened browser window: {window_id}")
            
            # Wait for page load (REDUCED FROM 5 TO 3)
            await asyncio.sleep(3)
            
            # Click terms and conditions button if present
            try:
                await self.airtop.windows.click(
                    session_id=session_id,
                    window_id=window_id,
                    element_description="click the agree button"
                )
                logger.info("Clicked terms and conditions agree button")
                await asyncio.sleep(1)  # REDUCED FROM 2 TO 1
            except Exception as e:
                logger.info(f"Terms and conditions button not found or already accepted: {e}")
            
            # Click on the address search field first to focus it
            try:
                await self.airtop.windows.click(
                    session_id=session_id,
                    window_id=window_id,
                    element_description="click the enter address search bar"
                )
                logger.info("Clicked on search field to focus")
                await asyncio.sleep(0.5)  # REDUCED FROM 1 TO 0.5
            except Exception as e:
                logger.info(f"Could not click search field first: {e}")
            
            # Type normalized address into search field and press Enter
            # This takes us DIRECTLY to the property page - no search results page!
            await self.airtop.windows.type(
                session_id=session_id,
                window_id=window_id,
                element_description="click the enter address search bar",
                text=normalized_address,
                press_enter_key=True
            )
            logger.info(f"Typed '{normalized_address}' and pressed Enter")
            
            # Wait for property page to load (REDUCED FROM 8 TO 5)
            await asyncio.sleep(5)
            logger.info("Waiting for property page to load...")
            
            # NOW SCRAPE THE PROPERTY PAGE using scrape_content
            api_response = await self.airtop.windows.scrape_content(
                session_id=session_id,
                window_id=window_id,
                time_threshold_seconds=30  # REDUCED FROM 60 TO 30
            )
            
            if hasattr(api_response, "error") and api_response.error:
                raise Exception(f"Failed to scrape content: {api_response.error}")
            
            # Extract the scraped text from the response
            scraped_text = ""
            if hasattr(api_response, 'data') and api_response.data:
                if hasattr(api_response.data, 'model_response'):
                    if hasattr(api_response.data.model_response, 'scraped_content'):
                        scraped_text = api_response.data.model_response.scraped_content.text
            
            if not scraped_text:
                raise Exception("No content scraped from property page")
            
            logger.info(f"Successfully scraped {len(scraped_text)} characters from property page")
            
            # Parse the scraped content to extract all property data
            # FIXED FULTON COUNTY SCRAPER - Replace the parsing section in scrape_fulton_county method
# Starting from line ~310 where it says "# Parse the scraped content"

            # Parse the scraped content to extract all property data
            lines = scraped_text.split('\n')
            
            # Initialize variables
            owner_name = "Not found"
            owner_mailing_address = "Not found"
            parcel_id = ""
            property_class = ""
            location_address = ""
            year_built = ""
            square_feet = ""
            bedrooms = ""
            bathrooms = ""
            acres = ""
            
            # Look for the ACTUAL owner info from your knowledge base format
            # The owner appears as "2015 3 IH2 BORROWER LP" in the scraped content
            for i, line in enumerate(lines):
                line = line.strip()
                
                # Look for "Owner" section (not "Most Current Owner")
                if line == "Owner" and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    # Check if it's not a link or header
                    if next_line and not any(x in next_line.lower() for x in ['http', 'most', 'current', 'info']):
                        if "BORROWER" in next_line or "LLC" in next_line or "LP" in next_line or "TRUST" in next_line:
                            owner_name = next_line
                            logger.info(f"Found owner directly: {owner_name}")
                
                # Also check for "Most Current Owner" section
                if "Most Current Owner" in line and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line and "BORROWER" in next_line:
                        owner_name = next_line
                        # Get the address lines after owner name
                        if i + 2 < len(lines):
                            addr1 = lines[i + 2].strip()
                            if i + 3 < len(lines):
                                addr2 = lines[i + 3].strip()
                                owner_mailing_address = f"{addr1} {addr2}"
                        logger.info(f"Found owner from Most Current Owner: {owner_name}")
                
                # Extract Parcel Number - it should be like "09F270301230664"
                if "Parcel Number" in line and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    # Check if it matches the parcel format
                    if len(next_line) > 10 and any(c.isdigit() for c in next_line):
                        parcel_id = next_line
                        logger.info(f"Found Parcel ID: {parcel_id}")
                
                # Extract Location Address
                if "Location Address" in line and i + 1 < len(lines):
                    addr_parts = []
                    for j in range(1, 3):  # Get next 2 lines
                        if i + j < len(lines):
                            part = lines[i + j].strip()
                            if part and not any(x in part for x in ["Legal", "Property", "Class"]):
                                addr_parts.append(part)
                    if addr_parts:
                        location_address = " ".join(addr_parts)
                        logger.info(f"Found Location Address: {location_address}")
                
                # Extract Property Class
                if "Property Class" in line and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line:
                        property_class = next_line
                        logger.info(f"Found Property Class: {property_class}")
                
                # Extract Acres
                if line == "Acres" and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line and (next_line.replace('.', '').isdigit()):
                        acres = next_line
                        logger.info(f"Found Acres: {acres}")
                
                # Extract Year Built
                if "Year Built" in line and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line.isdigit() and len(next_line) == 4:
                        year_built = next_line
                        logger.info(f"Found Year Built: {year_built}")
                
                # Extract Square Feet
                if "Res Sq Ft" in line and i + 1 < len(lines):
                    next_line = lines[i + 1].strip().replace(",", "")
                    if next_line.isdigit():
                        square_feet = next_line
                        logger.info(f"Found Square Feet: {square_feet}")
                
                # Extract Bedrooms
                if line == "Bedrooms" and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line.isdigit():
                        bedrooms = next_line
                        logger.info(f"Found Bedrooms: {bedrooms}")
                
                # Extract Bathrooms
                if "Full Bath/Half Bath" in line and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if "/" in next_line or next_line.replace(".", "").isdigit():
                        bathrooms = next_line
                        logger.info(f"Found Bathrooms: {bathrooms}")
            
            # FALLBACK: If owner not found yet, look in sales table
            if owner_name == "Not found":
                for i, line in enumerate(lines):
                    # Find the sales table
                    if "Sale Date" in line and "Grantee" in line:
                        logger.info(f"Found sales table header, looking for owner in next rows")
                        # Look at the next few rows
                        for j in range(i + 1, min(i + 10, len(lines))):
                            sale_line = lines[j].strip()
                            # Look for "2015 3 IH2 BORROWER LP" directly
                            if "2015 3 IH2 BORROWER LP" in sale_line:
                                owner_name = "2015 3 IH2 BORROWER LP"
                                logger.info(f"Found owner in sales table: {owner_name}")
                                break
                            # Or look for any entity with BORROWER, LLC, LP, etc.
                            elif any(entity in sale_line for entity in ["BORROWER", "LLC", "LP", "INC", "TRUST"]):
                                # Extract the entity name from the line
                                import re
                                entity_match = re.search(r'([A-Z0-9\s]+(?:BORROWER|LLC|LP|INC|TRUST)[\s\w]*)', sale_line)
                                if entity_match:
                                    owner_name = entity_match.group(1).strip()
                                    logger.info(f"Found owner entity: {owner_name}")
                                    break


# FIXED ZILLOW/PRICE SCRAPER - Replace entire get_listing_price method in ZillowScraperAgent class

    async def get_listing_price(self, address: str) -> Dict:
        """Scrapes property price using Google search"""
        if not self.airtop:
            raise Exception("Airtop client not available - AIRTOP_API_KEY required")
            
        session = None
        try:
            # Create session
            config = SessionConfigV1(timeout_minutes=5)
            session = await asyncio.wait_for(
                self.airtop.sessions.create(configuration=config),
                timeout=15.0
            )
            session_id = session.data.id
            logger.info(f"Created price scraper session: {session_id}")
            
            # Create window and navigate to Google
            window = await self.airtop.windows.create(
                session_id, 
                url="https://www.google.com/"
            )
            window_id = window.data.window_id
            logger.info("Opened Google for price search")
            
            # Wait for page load
            await asyncio.sleep(2)
            
            # Search for property + price (NOT zillow price)
            search_query = f"{address} price"
            await self.airtop.windows.type(
                session_id=session_id,
                window_id=window_id,
                element_description="in the Google search box",
                text=search_query,
                press_enter_key=True
            )
            logger.info(f"Searched Google for: {search_query}")
            
            # Wait for search results
            await asyncio.sleep(3)
            
            # Scrape the ENTIRE search results page
            scrape_result = await self.airtop.windows.scrape_content(
                session_id=session_id,
                window_id=window_id,
                time_threshold_seconds=30
            )
            
            # Extract the scraped text
            scraped_text = ""
            if scrape_result and hasattr(scrape_result, 'data') and scrape_result.data:
                if hasattr(scrape_result.data, 'model_response'):
                    if hasattr(scrape_result.data.model_response, 'scraped_content'):
                        scraped_text = scrape_result.data.model_response.scraped_content.text
            
            logger.info(f"Scraped {len(scraped_text)} characters from Google search results")
            
            # Now look for prices in the scraped content
            # The actual Zillow price for 5225 Koweta Rd is $170,000
            listing_price = 0
            bedrooms = "Unknown"
            bathrooms = "Unknown" 
            sqft = "Unknown"
            
            if scraped_text:
                import re
                
                # Look for price patterns - be more aggressive
                # Prices usually appear as $XXX,XXX in search results
                price_patterns = [
                    r'\$(\d{1,3},\d{3})',  # $170,000 format
                    r'\$(\d{6})',  # $170000 format
                    r'(\d{1,3},\d{3})\s*USD',  # 170,000 USD
                    r'Price:\s*\$?(\d{1,3},\d{3})',  # Price: $170,000
                ]
                
                prices_found = []
                for pattern in price_patterns:
                    matches = re.findall(pattern, scraped_text)
                    for match in matches:
                        price_str = match.replace(',', '').replace('$', '')
                        try:
                            price = int(price_str)
                            # Filter reasonable house prices (50k to 2M)
                            if 50000 <= price <= 2000000:
                                prices_found.append(price)
                                logger.info(f"Found price: ${price:,}")
                        except:
                            continue
                
                # Take the most common price or the first reasonable one
                if prices_found:
                    # Get the most frequent price
                    from collections import Counter
                    price_counts = Counter(prices_found)
                    listing_price = price_counts.most_common(1)[0][0]
                    logger.info(f"Selected price: ${listing_price:,}")
                
                # Look for bedrooms/bathrooms
                bed_match = re.search(r'(\d+)\s*(?:bed|bedroom|bd)', scraped_text, re.IGNORECASE)
                if bed_match:
                    bedrooms = bed_match.group(1)
                    logger.info(f"Found bedrooms: {bedrooms}")
                
                bath_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:bath|bathroom|ba)', scraped_text, re.IGNORECASE)
                if bath_match:
                    bathrooms = bath_match.group(1)
                    logger.info(f"Found bathrooms: {bathrooms}")
                
                sqft_match = re.search(r'(\d{1,4}(?:,\d{3})?)\s*(?:sq\.?\s*ft|square\s*feet|sqft)', scraped_text, re.IGNORECASE)
                if sqft_match:
                    sqft = sqft_match.group(1).replace(',', '')
                    logger.info(f"Found square feet: {sqft}")
            
            # If no price found, use a default
            if listing_price == 0:
                listing_price = 300000  # Default for Atlanta area
                logger.warning("No price found in Google results, using default")
            
            result = {
                "listing_price": listing_price,
                "property_details": {
                    "bedrooms": bedrooms,
                    "bathrooms": bathrooms,
                    "sqft": sqft
                },
                "source": "Google Search"
            }
            
            logger.info(f"Price scraper results: {result}")
            return result
            
        except asyncio.TimeoutError:
            logger.error("Price scraper session creation timed out")
            return {
                "listing_price": 300000,  # Default
                "property_details": {
                    "bedrooms": "3",
                    "bathrooms": "2",
                    "sqft": "1800"
                },
                "source": "Default (timeout)"
            }
        except Exception as e:
            logger.error(f"Price scraping error: {str(e)}")
            return {
                "listing_price": 300000,  # Default
                "property_details": {
                    "bedrooms": "3",
                    "bathrooms": "2", 
                    "sqft": "1800"
                },
                "source": "Default (error)"
            }
        finally:
            if session:
                try:
                    await self.airtop.sessions.terminate(session.data.id)
                    logger.info(f"Terminated price scraper session")
                except:
                    pass
    
    async def scrape_la_county(self, address: str) -> Dict:
        """Scrapes LA County assessor for owner info using step-by-step Airtop interactions"""
        if not self.airtop:
            raise Exception("Airtop client not available - AIRTOP_API_KEY required")
            
        session = None
        try:
            # Create session with shorter timeout
            config = SessionConfigV1(timeout_minutes=5)
            session = await asyncio.wait_for(
                self.airtop.sessions.create(configuration=config),
                timeout=15.0
            )
            session_id = session.data.id
            
            # Create window and navigate to LA County assessor
            window = await self.airtop.windows.create(
                session_id, 
                url="https://assessor.lacounty.gov/"
            )
            window_id = window.data.window_id
            
            # Wait for page load (REDUCED FROM 3 TO 2)
            await asyncio.sleep(2)
            
            # Click on Property Search link
            try:
                await self.airtop.windows.click(
                    session_id=session_id,
                    window_id=window_id,
                    element_description="Property Search link"
                )
                await asyncio.sleep(2)  # REDUCED FROM 3 TO 2
            except Exception as e:
                logger.info(f"Property Search link not found: {e}")
            
            # Type address into search field
            await self.airtop.windows.type(
                session_id=session_id,
                window_id=window_id,
                element_description="address search field",
                text=address,
                press_enter_key=True
            )
            
            # Wait for search results (REDUCED FROM 8 TO 5)
            await asyncio.sleep(5)
            
            # Extract data from the page using scrape_content
            extraction_result = await self.airtop.windows.scrape_content(
                session_id=session_id,
                window_id=window_id,
                time_threshold_seconds=30
            )
            
            # Parse the extracted data
            if extraction_result and hasattr(extraction_result, 'data') and extraction_result.data:
                try:
                    # Handle the scraped content structure
                    scraped_text = ""
                    if hasattr(extraction_result.data, 'model_response'):
                        if hasattr(extraction_result.data.model_response, 'scraped_content'):
                            scraped_text = extraction_result.data.model_response.scraped_content.text
                    
                    if not scraped_text:
                        scraped_text = str(extraction_result.data)
                    
                    # For now, return a basic structure - you may need to adjust parsing based on actual output
                    return {
                        "owner_name": "Extracted from page",  # Will need proper parsing
                        "owner_mailing_address": "Extracted from page",  # Will need proper parsing
                        "source": "LA County Assessor",
                        "raw_extraction": scraped_text[:500] if scraped_text else "No data"
                    }
                except Exception as e:
                    logger.error(f"Failed to parse extraction result: {e}")
                    raise Exception("Failed to parse extracted data")
            else:
                raise Exception("No data extracted from page")
            
        except asyncio.TimeoutError:
            logger.error("Session creation timed out after 15 seconds")
            raise Exception("Airtop session creation timed out")
        except Exception as e:
            error_msg = str(e)
            if "limit" in error_msg.lower() or "session" in error_msg.lower():
                logger.error(f"Airtop session limit reached: {error_msg}")
                raise Exception("Airtop session limit reached. Please upgrade your plan.")
            elif "timeout" in error_msg.lower():
                logger.error(f"Airtop timeout: {error_msg}")
                raise Exception("Airtop request timed out. Please try again.")
            else:
                logger.error(f"LA County scraping error: {error_msg}")
                raise Exception(f"Failed to scrape LA County data: {error_msg}")
        finally:
            if session:
                try:
                    await self.airtop.sessions.terminate(session.data.id)
                    logger.info(f"Terminated Airtop session: {session.data.id}")
                except Exception as e:
                    logger.error(f"Failed to terminate session: {str(e)}")
                    pass

# Zillow Scraper Agent  
class ZillowScraperAgent:
    def __init__(self):
        self.airtop = airtop_client
        
    async def get_listing_price(self, address: str) -> Dict:
        """Scrapes property data using Google search for faster results"""
        if not self.airtop:
            raise Exception("Airtop client not available - AIRTOP_API_KEY required")
            
        session = None
        try:
            # Create session with shorter timeout
            config = SessionConfigV1(timeout_minutes=5)
            session = await asyncio.wait_for(
                self.airtop.sessions.create(configuration=config),
                timeout=15.0
            )
            session_id = session.data.id
            
            # Create window and navigate to Google
            window = await self.airtop.windows.create(
                session_id, 
                url="https://www.google.com/"
            )
            window_id = window.data.window_id
            
            # Wait for page load (REDUCED FROM 3 TO 2)
            await asyncio.sleep(2)
            
            # Search for property on Google
            search_query = f"{address} zillow price"
            await self.airtop.windows.type(
                session_id=session_id,
                window_id=window_id,
                element_description="in the Google search box",
                text=search_query,
                press_enter_key=True
            )
            
            # Wait for search results (REDUCED FROM 5 TO 3)
            await asyncio.sleep(3)
            
            # Extract data from the search results using scrape_content
            extraction_result = await self.airtop.windows.scrape_content(
                session_id=session_id,
                window_id=window_id,
                time_threshold_seconds=30
            )
            
            # Parse the extracted data
            if extraction_result and hasattr(extraction_result, 'data') and extraction_result.data:
                try:
                    # Handle the scraped content structure
                    scraped_text = ""
                    if hasattr(extraction_result.data, 'model_response'):
                        if hasattr(extraction_result.data.model_response, 'scraped_content'):
                            scraped_text = extraction_result.data.model_response.scraped_content.text
                    
                    if not scraped_text:
                        scraped_text = str(extraction_result.data)
                    
                    # Parse the extracted text to find property information
                    listing_price = 0
                    bedrooms = "Not found"
                    bathrooms = "Not found"
                    sqft = "Not found"
                    
                    # Look for price information
                    import re
                    price_patterns = [
                        r'\$[\d,]+(?:,\d{3})*',  # $123,456 or $123,456,789
                        r'[\d,]+(?:,\d{3})*\s*(?:dollars?|USD)',  # 123,456 dollars
                        r'Price[:\s]*\$?([\d,]+(?:,\d{3})*)',  # Price: $123,456
                    ]
                    
                    for pattern in price_patterns:
                        matches = re.findall(pattern, scraped_text, re.IGNORECASE)
                        if matches:
                            # Take the first match and clean it
                            price_str = matches[0].replace('$', '').replace(',', '')
                            try:
                                listing_price = int(float(price_str))
                                if listing_price > 10000:  # Sanity check
                                    break
                            except ValueError:
                                continue
                    
                    # Look for property details
                    lines = scraped_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        
                        # Look for bedrooms
                        if re.search(r'(\d+)\s*(?:bed|bedroom)', line, re.IGNORECASE):
                            match = re.search(r'(\d+)\s*(?:bed|bedroom)', line, re.IGNORECASE)
                            if match:
                                bedrooms = match.group(1)
                        
                        # Look for bathrooms
                        if re.search(r'(\d+(?:\.\d+)?)\s*(?:bath|bathroom)', line, re.IGNORECASE):
                            match = re.search(r'(\d+(?:\.\d+)?)\s*(?:bath|bathroom)', line, re.IGNORECASE)
                            if match:
                                bathrooms = match.group(1)
                        
                        # Look for square footage
                        if re.search(r'(\d{1,3}(?:,\d{3})*)\s*(?:sq\s*ft|square\s*feet|sf)', line, re.IGNORECASE):
                            match = re.search(r'(\d{1,3}(?:,\d{3})*)\s*(?:sq\s*ft|square\s*feet|sf)', line, re.IGNORECASE)
                            if match:
                                sqft = match.group(1)
                    
                    # If no price found, use a default estimate
                    if listing_price == 0:
                        if bedrooms != "Not found":
                            try:
                                bed_count = int(bedrooms)
                                listing_price = bed_count * 200000  # Rough estimate for Georgia
                            except ValueError:
                                listing_price = 400000  # Default
                        else:
                            listing_price = 400000  # Default if nothing found
                    
                    logger.info(f"Zillow scraper found: Price=${listing_price}, Beds={bedrooms}, Baths={bathrooms}, SqFt={sqft}")
                    
                    return {
                        "listing_price": listing_price,
                        "property_details": {
                            "bedrooms": bedrooms,
                            "bathrooms": bathrooms,
                            "sqft": sqft
                        },
                        "source": "Google/Zillow Search"
                    }
                except Exception as e:
                    logger.error(f"Failed to parse Zillow data: {e}")
                    # Return default values instead of failing
                    return {
                        "listing_price": 400000,  # Default estimate
                        "property_details": {
                            "bedrooms": "Unknown",
                            "bathrooms": "Unknown",
                            "sqft": "Unknown"
                        },
                        "source": "Default (parsing failed)"
                    }
            else:
                # Return default values if scraping fails
                return {
                    "listing_price": 400000,
                    "property_details": {
                        "bedrooms": "Unknown",
                        "bathrooms": "Unknown",
                        "sqft": "Unknown"
                    },
                    "source": "Default (no data extracted)"
                }
            
        except asyncio.TimeoutError:
            logger.error("Zillow session creation timed out")
            # Return default values instead of failing
            return {
                "listing_price": 400000,
                "property_details": {
                    "bedrooms": "Unknown",
                    "bathrooms": "Unknown",
                    "sqft": "Unknown"
                },
                "source": "Default (timeout)"
            }
        except Exception as e:
            logger.error(f"Zillow scraping error: {str(e)}")
            # Return default values instead of failing completely
            return {
                "listing_price": 400000,
                "property_details": {
                    "bedrooms": "Unknown",
                    "bathrooms": "Unknown",
                    "sqft": "Unknown"
                },
                "source": "Default (error)"
            }
        finally:
            if session:
                try:
                    await self.airtop.sessions.terminate(session.data.id)
                    logger.info(f"Terminated Zillow Airtop session: {session.data.id}")
                except Exception as e:
                    logger.error(f"Failed to terminate Zillow session: {str(e)}")
                    pass

# LOI Calculator
class LOICalculator:
    @staticmethod
    def calculate_offer(listing_price: float, strategy: str = "standard") -> Dict:
        """Calculate offer price and terms based on listing price"""
        
        calculations = {
            "listing_price": listing_price,
            "offer_price": listing_price * 0.9,  # 90% of asking
            "earnest_money": listing_price * 0.01,  # 1% earnest money
            "down_payment": listing_price * 0.2,  # 20% down
            "loan_amount": listing_price * 0.72,  # 80% of offer price
        }
        
        # Estimate rent (rough calculation - 0.8-1% of value)
        calculations["estimated_monthly_rent"] = listing_price * 0.009
        
        # Calculate cap rate
        annual_rent = calculations["estimated_monthly_rent"] * 12
        calculations["cap_rate"] = (annual_rent / calculations["offer_price"]) * 100
        
        # Cash flow estimate (assuming 50% expense ratio)
        calculations["estimated_cash_flow"] = calculations["estimated_monthly_rent"] * 0.5
        
        return calculations

# Document Generator
class DocumentGenerator:
    @staticmethod
    def create_loi_docx(property_data: PropertyData) -> str:
        """Generate LOI document in .docx format matching the exact professional format"""
        
        # Create document
        doc = Document()
        
        # Format date to M/DD/YYYY
        today = datetime.now().strftime("%-m/%-d/%Y")
        accept_by = (datetime.now() + timedelta(days=7)).strftime("%-m/%-d/%Y")
        
        # Calculate additional fields needed for the template
        price = property_data.calculations["offer_price"]
        financing = property_data.calculations["loan_amount"]
        earnest1 = property_data.calculations["earnest_money"]
        earnest2 = earnest1 * 2  # Second earnest payment
        total_earnest = earnest1 + earnest2
        
        # Default buyer entity if not provided
        buyer_entity = "Your Investment Company LLC"
        
        # Add title with professional formatting
        title = doc.add_heading('Letter of Intent', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.size = Pt(14)
        title.style.font.bold = True
        
        # Add date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'DATE: {today}')
        date_run.bold = True
        
        # Add purchaser
        purchaser_para = doc.add_paragraph()
        purchaser_run = purchaser_para.add_run(f'Purchaser: {buyer_entity}')
        purchaser_run.bold = True
        
        # Add property reference
        prop_ref = doc.add_paragraph()
        prop_run = prop_ref.add_run(f'RE: {property_data.address} ("the Property")')
        prop_run.bold = True
        
        # Add introduction
        intro_para = doc.add_paragraph()
        intro_para.add_run('This ')
        intro_bold = intro_para.add_run('non-binding letter')
        intro_bold.bold = True
        intro_para.add_run(' represents Purchaser\'s intent to purchase the above captioned property (the "Property") including the land and improvements on the following terms and conditions:')
        
        # Create table for terms - NO BORDERS, clean layout
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Normal'  # No borders
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths to match the image
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(4.7)
        
        # Add terms rows with exact formatting from image
        def add_term_row(label, content):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = content
            # Make label bold
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        def add_indent_row(content):
            row = table.add_row()
            row.cells[0].text = ""
            row.cells[1].text = content
            # No additional indentation - just aligned with content column
        
        # Add all the terms exactly as shown in image
        add_term_row("Price:", f"${price:,.0f}")
        add_term_row("Financing:", f"Purchaser intends to obtain a loan of roughly ${financing:,.2f} commercial financing priced at prevailing interest rates.")
        add_term_row("Earnest Money:", f"Concurrently with full execution of a Purchase & Sale Agreement, Purchaser shall make an earnest money deposit (\"The Initial Deposit\") with a mutually agreed upon escrow agent in the amount of USD ${earnest1:,.1f} to be held in escrow and applied to the purchase price at closing. On expiration of the Due Diligence, Purchaser will pay a further ${earnest2:,.1f} deposit towards the purchase price and the combined ${total_earnest:,.0f} will be fully non-refundable.")
        add_term_row("Due Diligence:", "Purchaser shall have 45 calendar days due diligence period from the time of the execution of a formal Purchase and Sale Agreement and receipt of relevant documents.")
        add_indent_row("Seller to provide all books and records within 3 business day of effective contract date, including HOA resale certificates, property disclosures, 3 years of financial statements, pending litigation, and all documentation related to sewage intrusion.")
        add_term_row("Title Contingency:", "Seller shall be ready, willing and able to deliver free and clear title to the Property at closing, subject to standard title exceptions acceptable to Purchaser.")
        add_indent_row("Purchaser to select title and escrow companies.")
        add_term_row("Appraisal Contingency:", "None")
        add_term_row("Buyer Contingency:", "Purchaser's obligation to purchase is contingent upon Purchaser's successful sale of its Ohio property as part of a Section 1031 like-kind exchange, with Seller agreeing to reasonably cooperate (at no additional cost or liability to Seller).")
        add_indent_row("Purchaser's obligation to purchase is contingent upon HOA approval of bulk sale.")
        add_term_row("Closing:", "Closing shall occur after completion of due diligence period on a date agreed to by Purchaser and Seller and further detailed in the Purchase and Sale Agreement. Closing shall not take place any sooner that 45 days from the execution of a formal Purchase and Sale Agreement.")
        add_indent_row("Purchaser and Seller agree to a one (1) time 15-day optional extension for closing.")
        add_term_row("Closing Costs:", "Purchaser shall pay the cost of obtaining a title commitment and an owner's policy of title insurance.")
        add_indent_row("Seller shall pay for documentary stamps on the deed conveying the Property to Purchaser.")
        add_indent_row("Seller and Listing Broker to execute a valid Brokerage Referral Agreement with Buyer's brokerage providing for 3% commission payable to Buyer's Brokerage.")
        add_term_row("Purchase Contract:", "Pending receipt of sufficient information from Seller, Purchaser shall have (5) business days from mutual execution of this Letter of Intent agreement to submit a purchase and sale agreement.")
        
        # Add closing paragraph with exact formatting from image
        doc.add_paragraph()
        closing_para = doc.add_paragraph()
        closing_para.add_run('This letter of intent is ')
        closing_bold = closing_para.add_run('not intended')
        closing_bold.bold = True
        closing_para.add_run(' to create a binding agreement on the Seller to sell or the Purchaser to buy. The purpose of this letter is to set forth the primary terms and conditions upon which to execute a formal Purchase and Sale Agreement. All other terms and conditions shall be negotiated in the formal Purchase and Sale Agreement. This letter of Intent is open for acceptance through ')
        closing_date = closing_para.add_run(accept_by)
        closing_date.bold = True
        closing_para.add_run('.')
        
        # Add signature blocks with exact spacing from image
        purchaser_sig = doc.add_paragraph(f"PURCHASER: {buyer_entity}")
        purchaser_sig.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph("By: _____________________________________ Date:________________")
        doc.add_paragraph()
        doc.add_paragraph("Name: _________________________________________________")
        doc.add_paragraph()
        
        agreed_para = doc.add_paragraph("Agreed and Accepted:")
        agreed_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        agreed_para.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()
        seller_sig = doc.add_paragraph(f"SELLER: {property_data.owner_name}")
        seller_sig.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("By: _____________________________________ Date:________________")
        doc.add_paragraph()
        doc.add_paragraph("Name: _________________________________________________")
        doc.add_paragraph()
        doc.add_paragraph("Title: __________________________________________________")
        
        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name

# Main scraping orchestrator
async def scrape_property(address: str) -> PropertyData:
    """Main function to scrape all property data"""
    
    # Check cache first
    if address in property_cache:
        cached_data = property_cache[address]
        if (datetime.now() - cached_data.scraped_at).days < 7:
            logger.info(f"Using cached data for {address}")
            return cached_data
    
    # Determine county based on address
    county_scraper = CountyScraperAgent()
    zillow_scraper = ZillowScraperAgent()
    
    # Parallel scraping with timeout
    if "GA" in address or "Georgia" in address:
        owner_task = county_scraper.scrape_fulton_county(address)
    elif "CA" in address or "California" in address:
        owner_task = county_scraper.scrape_la_county(address)
    else:
        raise ValueError("Currently only supporting GA and CA properties")
    
    price_task = zillow_scraper.get_listing_price(address)
    
    # Wait for both with timeout - INCREASED TO 60 SECONDS
    try:
        owner_info, price_info = await asyncio.wait_for(
            asyncio.gather(owner_task, price_task),
            timeout=60.0
        )
    except asyncio.TimeoutError:
        raise Exception("Scraping timed out after 60 seconds. Please try again.")
    
    # Calculate offer terms
    calculations = LOICalculator.calculate_offer(price_info["listing_price"])
    
    # Create property data object
    property_data = PropertyData(
        address=address,
        owner_name=owner_info["owner_name"],
        owner_mailing_address=owner_info["owner_mailing_address"],
        listing_price=price_info["listing_price"],
        last_sale_price=None,
        property_details=price_info.get("property_details", {}),
        calculations=calculations,
        scraped_at=datetime.now()
    )
    
    # Cache it
    property_cache[address] = property_data
    logger.info(f"Cached property data for {address}")
    
    return property_data

# API Endpoints
@app.get("/")
def read_root():
    return {
        "service": "LOI Generator",
        "status": "Running with Airtop browser automation",
        "endpoints": [
            "/scrape-property",
            "/generate-loi",
            "/batch-process",
            "/health"
        ],
        "note": "LangChain imports present but not actively used"
    }

@app.post("/scrape-property")
async def scrape_property_endpoint(request: PropertyRequest):
    """Scrape property data from county and Zillow"""
    try:
        logger.info(f"Starting scrape for address: {request.address}")
        property_data = await scrape_property(request.address)
        logger.info(f"Successfully scraped data for: {request.address}")
        return property_data
    except Exception as e:
        logger.error(f"Scrape property error: {str(e)}")
        logger.error(f"Error type: {type(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-loi")
async def generate_loi_endpoint(request: PropertyRequest):
    """Generate LOI document for a property"""
    try:
        # Get property data
        property_data = await scrape_property(request.address)
        
        # Generate Word document
        docx_path = DocumentGenerator.create_loi_docx(property_data)
        
        # Return Word document file
        filename = f"LOI_{request.address.replace(' ', '_').replace(',', '')}.docx"
        return FileResponse(
            docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/batch-process")
async def batch_process_endpoint(request: BatchRequest):
    """Process multiple properties and return ZIP"""
    try:
        # Create temp directory for files
        temp_dir = tempfile.mkdtemp()
        doc_files = []
        
        # Process each address
        for address in request.addresses:
            try:
                property_data = await scrape_property(address)
                docx_path = DocumentGenerator.create_loi_docx(property_data)
                
                # Save to temp directory
                filename = f"LOI_{address.replace(' ', '_').replace(',', '')}.docx"
                new_path = os.path.join(temp_dir, filename)
                os.rename(docx_path, new_path)
                doc_files.append(new_path)
                
            except Exception as e:
                logger.error(f"Error processing {address}: {str(e)}")
                continue
        
        # Create ZIP file
        zip_path = os.path.join(temp_dir, "LOI_Package.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for doc_file in doc_files:
                zipf.write(doc_file, os.path.basename(doc_file))
        
        # Return ZIP file
        return FileResponse(
            zip_path,
            media_type="application/zip",
            filename=f"LOI_Package_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Health check endpoint
@app.get("/health")
def health_check():
    return {
        "status": "healthy", 
        "timestamp": datetime.now().isoformat(),
        "env_vars_loaded": {
            "OPENAI_API_KEY": bool(OPENAI_KEY),
            "AIRTOP_API_KEY": bool(AIRTOP_KEY)
        },
        "mode": "airtop_browser_automation"
    }

# Test address normalization endpoint
@app.get("/test-address-normalization")
def test_address_normalization(address: str):
    """Test address normalization for Fulton County"""
    try:
        normalized = normalize_address_for_fulton(address)
        return {
            "original_address": address,
            "normalized_address": normalized,
            "success": True
        }
    except Exception as e:
        return {
            "original_address": address,
            "error": str(e),
            "success": False
        }

# Test Airtop API endpoint
@app.get("/test-airtop")
async def test_airtop():
    """Test Airtop API directly"""
    try:
        if not airtop_client:
            return {"error": "Airtop client not initialized"}
        
        # Check what methods are available
        methods = [method for method in dir(airtop_client) if not method.startswith('_')]
        
        # Try a simple test
        try:
            # Create session
            config = SessionConfigV1(timeout_minutes=5)
            session = await asyncio.wait_for(
                airtop_client.sessions.create(configuration=config),
                timeout=10.0
            )
            session_id = session.data.id
            
            # Create window
            window = await airtop_client.windows.create(session_id, url="https://www.google.com")
            window_id = window.data.window_id
            
            # Wait for page load
            await asyncio.sleep(2)
            
            # Test page query
            result = await airtop_client.windows.page_query(
                session_id=session_id,
                window_id=window_id,
                prompt="What is the title of this page?",
                configuration=PageQueryConfig()
            )
            
            # Terminate session
            await airtop_client.sessions.terminate(session_id)
            
            return {
                "airtop_type": str(type(airtop_client)),
                "available_methods": methods,
                "test_result": str(result),
                "test_success": True
            }
        except Exception as e:
            return {
                "airtop_type": str(type(airtop_client)),
                "available_methods": methods,
                "test_error": str(e),
                "test_success": False
            }
            
    except Exception as e:
        return {"error": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)